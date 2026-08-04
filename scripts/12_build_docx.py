"""N2 step 12 - render the Frontiers submission set as .docx.

pandoc and node are unavailable on this machine; python-docx + LibreOffice are.
Produces: manuscript (with continuous line numbers, as Frontiers requires),
cover letter, and supplementary material.
"""
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(r"D:\N2_icu_nutrition_delivery_gap")
MAN, SUB = ROOT / "07_manuscript", ROOT / "08_submission"
SUB.mkdir(exist_ok=True)

FONT, SIZE = "Times New Roman", Pt(12)


def _add_page_numbers(sec):
    """Centered 'Page N' field in the footer."""
    par = sec.footer.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run("Page ")
    for instr, kind in (("begin", None), (None, "PAGE"), ("end", None)):
        r = par.add_run()._r
        if kind:
            fld = OxmlElement("w:instrText")
            fld.set(qn("xml:space"), "preserve")
            fld.text = " PAGE "
            r.append(fld)
        else:
            f = OxmlElement("w:fldChar")
            f.set(qn("w:fldCharType"), instr)
            r.append(f)


def base_doc(line_numbers=False):
    d = Document()
    st = d.styles["Normal"]
    st.font.name = FONT
    st.font.size = SIZE
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = st.paragraph_format
    pf.line_spacing = 1.0
    pf.space_after = Pt(0)
    for s in d.sections:
        s.top_margin = s.bottom_margin = Inches(1)
        s.left_margin = s.right_margin = Inches(1)
        _add_page_numbers(s)
        if line_numbers:
            ln = OxmlElement("w:lnNumType")
            ln.set(qn("w:countBy"), "1")
            ln.set(qn("w:restart"), "continuous")
            ln.set(qn("w:distance"), "360")
            s._sectPr.append(ln)
    return d


INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*\n]+?\*|`[^`\n]+?`)", re.S)


def add_runs(par, text):
    """Render **bold**, *italic*, `code` inline."""
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = par.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            r = par.add_run(tok[1:-1]); r.italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = par.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9)
        else:
            par.add_run(tok)


def heading(d, text, level):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14 if level == 1 else 12)
    r.font.color.rgb = RGBColor(0, 0, 0)
    return p


def md_table(d, rows):
    hdr = [c.strip() for c in rows[0].strip().strip("|").split("|")]
    body = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows[2:]]
    t = d.add_table(rows=1, cols=len(hdr))
    t.style = "Table Grid"
    for i, h in enumerate(hdr):
        c = t.rows[0].cells[i]
        c.text = ""
        add_runs(c.paragraphs[0], h)
        for run in c.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
    for br in body:
        cells = t.add_row().cells
        for i, v in enumerate(br[:len(hdr)]):
            cells[i].text = ""
            add_runs(cells[i].paragraphs[0], v)
            for run in cells[i].paragraphs[0].runs:
                run.font.size = Pt(8)
    for row in t.rows:
        for c in row.cells:
            c.paragraphs[0].paragraph_format.line_spacing = 1.0
            c.paragraphs[0].paragraph_format.space_after = Pt(0)
    d.add_paragraph()


def convert(md_text, d, skip_sections=()):
    lines = md_text.split("\n")
    i, buf, skipping = 0, [], False

    def flush():
        nonlocal buf
        if buf:
            txt = " ".join(x.strip() for x in buf).strip()
            if txt:
                p = d.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                add_runs(p, txt)
            buf = []

    while i < len(lines):
        ln = lines[i]
        if ln.startswith("#"):
            flush()
            lvl = len(ln) - len(ln.lstrip("#"))
            title = ln.lstrip("#").strip()
            skipping = any(s.lower() in title.lower() for s in skip_sections)
            if not skipping:
                heading(d, title, 1 if lvl <= 2 else 2)
            i += 1
            continue
        if skipping:
            i += 1
            continue
        if ln.strip().startswith("|") and i + 1 < len(lines) and set(
                lines[i + 1].replace("|", "").replace(":", "").strip()) <= {"-", " "}:
            flush()
            tbl = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl.append(lines[i]); i += 1
            md_table(d, tbl)
            continue
        if ln.strip().startswith("```"):
            flush()
            i += 1
            code = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i]); i += 1
            i += 1
            for c in code:
                p = d.add_paragraph()
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = Pt(0)
                r = p.add_run(c)
                r.font.name = "Consolas"; r.font.size = Pt(8)
            d.add_paragraph()
            continue
        if ln.strip() in ("---", "***"):
            flush(); i += 1; continue
        if ln.strip().startswith("- "):
            flush()
            p = d.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            add_runs(p, ln.strip()[2:])
            i += 1
            continue
        if not ln.strip():
            flush()
        else:
            buf.append(ln)
        i += 1
    flush()


# =============================================================== manuscript
src = (MAN / "FrontNutr_main_manuscript.md").read_text(encoding="utf-8")
title = src.split("\n", 1)[0].lstrip("# ").strip()
body = src.split("\n", 1)[1]
body = re.sub(r"^\*\*Target journal:\*\*.*?\n", "", body, flags=re.M)

d = base_doc(line_numbers=True)

# (name, superscript) pairs - real superscript runs, not Unicode look-alikes
AUTHOR_PARTS = [("Jiajun Luo", "1,2,†"), ("Qinglong Chen", "3,†"),
                ("Jing Liu", "3,†"), ("Fanghui Lu", "3,*"), ("Xiaolong Liang", "1,*")]
AFFIL = [
    "1  Department of Gastrointestinal Surgery, The First Affiliated Hospital of "
    "Chongqing Medical University, Chongqing, China",
    "2  Molecular Oncology Laboratory, Department of Orthopedic Surgery and "
    "Rehabilitation Medicine, The University of Chicago Medical Center, Chicago, IL, USA",
    "3  Department of Cancer Center, The Second Affiliated Hospital of Chongqing "
    "Medical University, Chongqing, China",
]
CORR = [
    "Xiaolong Liang, MD, PhD — Department of Gastrointestinal Surgery, The First "
    "Affiliated Hospital of Chongqing Medical University, Chongqing 400016, China. "
    "E-mail: 204951@hospital.cqmu.edu.cn",
    "Fanghui Lu, PhD — Department of Cancer Center, The Second Affiliated Hospital "
    "of Chongqing Medical University, Chongqing, China. E-mail: lufh@cqmu.edu.cn",
]

p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(14)
r = p.add_run(title); r.bold = True; r.font.size = Pt(16)

p = d.add_paragraph(); p.paragraph_format.space_after = Pt(10)
for i, (nm, sup) in enumerate(AUTHOR_PARTS):
    if i:
        p.add_run(", ")
    p.add_run(nm)
    rs = p.add_run(sup); rs.font.superscript = True

for a in AFFIL:
    p = d.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(a); r.font.size = Pt(10)

for note in ("† These authors have contributed equally to this work and share first "
             "authorship.", "* Correspondence:"):
    p = d.add_paragraph(); p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.line_spacing = 1.15
    r = p.add_run(note); r.font.size = Pt(10)

for c in CORR:
    p = d.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run(c); r.font.size = Pt(10)

import re as _re
_body_txt = body.split("## 1 Introduction", 1)[1].split("## References", 1)[0]
_clean = _re.sub(r"[*`]", "", _re.sub(r"\{\{[^}]*\}\}", "", _body_txt))
_words = sum(1 for w in _clean.split() if _re.search(r"[A-Za-z0-9]", w))  # Word-equivalent
_nfig = len(_re.findall(r"^\*\*Figure \d", body, _re.M))
_ntab = len(_re.findall(r"^\*\*Table \d", body, _re.M))

p = d.add_paragraph(); p.paragraph_format.space_before = Pt(12)
rr = p.add_run("Article type: "); rr.bold = True
p.add_run("Original Research — Frontiers in Nutrition, Clinical Nutrition section")
for lbl, val in (("Word count (Introduction to Conclusion): ", f"{_words:,}"),
                 ("Figures: ", str(_nfig)), ("Tables: ", str(_ntab)),
                 ("Supplementary files: ", "1")):
    p = d.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    rr = p.add_run(lbl); rr.bold = True; rr.font.size = Pt(10)
    r2 = p.add_run(val); r2.font.size = Pt(10)

d.add_page_break()
convert(body, d, skip_sections=("Author-side items outstanding",))
out = SUB / "N2_FrontNutr_manuscript.docx"
d.save(out)
print(f"  {out.name}")

# =============================================================== cover letter
d = base_doc()
for _s in d.sections:                       # tighter page so the letter fits one page
    _s.top_margin = _s.bottom_margin = Inches(0.7)
    _s.left_margin = _s.right_margin = Inches(0.9)
d.styles["Normal"].font.size = Pt(11)
convert((MAN / "cover_letter_frontnutr.md").read_text(encoding="utf-8"), d)
d.save(SUB / "N2_cover_letter.docx")
print("  N2_cover_letter.docx")

# =============================================================== supplement
d = base_doc()
sup = (MAN / "supplement.md").read_text(encoding="utf-8")
sup = re.sub(r"^# .*$", "", sup, flags=re.M)          # no repeated title
sup = re.sub(r"^\*\*(Chance|Background) co-occurrence.*$", "", sup, flags=re.M)
convert(sup, d)
d.save(SUB / "N2_Supplementary_Material.docx")
print("  N2_Supplementary_Material.docx")

print("\nDONE")
