"""N2 step 52 - render and assemble the BMJ Quality & Safety submission package."""
import hashlib
import json
import re
import shutil
import zipfile
from datetime import datetime, timezone
from pathlib import Path

import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(r"D:\N2_icu_nutrition_delivery_gap")
MAN, SUB, FIG, OUT = (ROOT / "07_manuscript", ROOT / "08_submission",
                      ROOT / "04_figures", ROOT / "03_outputs")
REV, CAN = OUT / "review6", OUT / "canonical"
SUB.mkdir(exist_ok=True)

# reuse the converter from script 12 rather than reimplementing it
_src = (ROOT / "01_scripts" / "12_build_docx.py").read_text(encoding="utf-8")
_ns = {}
exec(compile(_src.split("# ========================================"
                        "======================= manuscript")[0],
             "docx_helpers", "exec"), _ns)
base_doc, convert = _ns["base_doc"], _ns["convert"]

AUTHORS = [("Jiajun Luo", "1,2,\u2020"), ("Qinglong Chen", "3,\u2020"),
           ("Jing Liu", "3,\u2020"), ("Fanghui Lu", "3,*"), ("Xiaolong Liang", "1,*")]
AFFIL = [
    "1  Department of Gastrointestinal Surgery, The First Affiliated Hospital of "
    "Chongqing Medical University, Chongqing, China",
    "2  Molecular Oncology Laboratory, Department of Orthopedic Surgery and "
    "Rehabilitation Medicine, The University of Chicago Medical Center, Chicago, IL, USA",
    "3  Department of Cancer Center, The Second Affiliated Hospital of Chongqing "
    "Medical University, Chongqing, China",
]
CORR = [
    "Xiaolong Liang, MD, PhD \u2014 Department of Gastrointestinal Surgery, The First "
    "Affiliated Hospital of Chongqing Medical University, Chongqing 400016, China. "
    "204951@hospital.cqmu.edu.cn",
    "Fanghui Lu, PhD \u2014 Department of Cancer Center, The Second Affiliated Hospital "
    "of Chongqing Medical University, Chongqing, China. lufh@cqmu.edu.cn",
]


def wc(s):
    s = "\n".join(l for l in s.split("\n")
                  if not l.lstrip().startswith("|") and not l.startswith("**Table "))
    return sum(1 for w in re.sub(r"[*`]", "", s).split() if re.search(r"[A-Za-z0-9]", w))


# =============================================================== manuscript
src = (MAN / "BMJQS_main_manuscript.md").read_text(encoding="utf-8")
title = src.split("\n", 1)[0].lstrip("# ").strip()
body = src.split("\n", 1)[1].split("<!-- generated back matter -->")[0]

d = base_doc(line_numbers=True)
p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(14)
r = p.add_run(title)
r.bold = True
r.font.size = Pt(15)

p = d.add_paragraph()
p.paragraph_format.space_after = Pt(10)
for i, (nm, sup) in enumerate(AUTHORS):
    if i:
        p.add_run(", ")
    p.add_run(nm)
    p.add_run(sup).font.superscript = True
for a in AFFIL:
    q = d.add_paragraph()
    q.paragraph_format.space_after = Pt(2)
    q.paragraph_format.line_spacing = 1.15
    q.add_run(a).font.size = Pt(10)
for note in ("\u2020 These authors contributed equally and share first authorship.",
             "* Joint corresponding authors and guarantors:"):
    q = d.add_paragraph()
    q.paragraph_format.space_before = Pt(6)
    q.paragraph_format.space_after = Pt(2)
    q.add_run(note).font.size = Pt(10)
for c in CORR:
    q = d.add_paragraph()
    q.paragraph_format.space_after = Pt(2)
    q.paragraph_format.left_indent = Inches(0.25)
    q.add_run(c).font.size = Pt(10)

_ab = body.split("## Abstract", 1)[1].split("**Keywords:**", 1)[0]
_bd = body.split("## Introduction", 1)[1].split("## Abbreviations", 1)[0]
for lbl, val in (("Article type: ", "Original Research"),
                 ("Abstract word count: ", f"{wc(_ab):,}"),
                 ("Main text word count (Introduction to Conclusion): ", f"{wc(_bd):,}"),
                 ("Figures: ", "3"), ("Tables: ", "2"), ("Additional files: ", "1")):
    q = d.add_paragraph()
    q.paragraph_format.space_after = Pt(2)
    rr = q.add_run(lbl)
    rr.bold = True
    rr.font.size = Pt(10)
    q.add_run(val).font.size = Pt(10)
d.add_page_break()
convert(body, d)
d.save(SUB / "BMJQS_manuscript.docx")
print(f"  BMJQS_manuscript.docx ({len(d.tables)} tables)")

# =============================================================== cover letter
d = base_doc()
for s in d.sections:
    s.top_margin = s.bottom_margin = Inches(0.7)
    s.left_margin = s.right_margin = Inches(0.9)
d.styles["Normal"].font.size = Pt(11)
convert((MAN / "cover_letter_bmjqs.md").read_text(encoding="utf-8"), d)
d.save(SUB / "BMJQS_cover_letter.docx")
print("  BMJQS_cover_letter.docx")

# =============================================================== additional file
d = base_doc()
af_md = (MAN / "BMJQS_additional_file.md").read_text(encoding="utf-8")
convert(re.sub(r"^# .*$", "", af_md, flags=re.M), d)
d.save(SUB / "BMJQS_Additional_file_1.docx")
print(f"  BMJQS_Additional_file_1.docx ({len(d.tables)} tables)")

# =============================================================== package + QA
PKG = SUB / "bmjqs_package"
if PKG.exists():
    shutil.rmtree(PKG)
PKG.mkdir()
for i, f in enumerate(["figure1_cohort_flow", "figure3_attribution",
                       "figure3_comparability"], 1):
    for ext in ("tif", "pdf"):
        if (FIG / f"{f}.{ext}").exists():
            shutil.copy(FIG / f"{f}.{ext}", PKG / f"Figure{i}.{ext}")
for a, b in [("BMJQS_manuscript.docx", "Manuscript.docx"),
             ("BMJQS_cover_letter.docx", "Cover_Letter.docx"),
             ("BMJQS_Additional_file_1.docx", "Additional_file_1.docx")]:
    shutil.copy(SUB / a, PKG / b)
for n in ("references_verified.csv", "bmjqs_validation.csv"):
    shutil.copy(OUT / n, PKG / n)

ms = (MAN / "BMJQS_main_manuscript.md").read_text(encoding="utf-8")
MSF = re.sub(r"\s+", " ", ms)
af = af_md
cov = (MAN / "cover_letter_bmjqs.md").read_text(encoding="utf-8")
val = pd.read_csv(OUT / "bmjqs_validation.csv")
reg = pd.read_csv(OUT / "exploratory_attempts.csv")
reg_ids = sorted(reg["attempt"].astype(str))
SCL = pd.read_csv(REV / "scale_and_priority_free.csv")
ref_pct = SCL[SCL.numerator_scale.str.startswith("reference")].iloc[0].pct_of_shortfall

checks = [
    ("all manuscript claims validated", int(val["pass"].sum()) == len(val),
     f"{int(val['pass'].sum())}/{len(val)}"),
    ("STROBE checklist supplied", "## S0. STROBE checklist" in af, "S0"),
    ("abbreviations list present", "## Abbreviations" in ms, "yes"),
    ("AI use stated in Methods", "### Use of generative AI" in ms.split("## Results")[0],
     "in Methods"),
    ("both plans hashed in appendix",
     "307a6452" in af and "ca02b0d6" in af, "primary + external"),
    ("registry complete in appendix", all(r in af for r in reg_ids),
     f"{len(reg_ids)} entries E01-{reg_ids[-1]}"),
    ("replication claim withdrawn everywhere",
     all("Same answer" not in x for x in (ms, af, cov))
     and "no claim of replication" in MSF and "make no claim of external replication" in
     re.sub(r"\s+", " ", cov), "withdrawn in text, appendix and letter"),
    ("like-for-like airway result present", "0.74" in MSF and "0.74" in af, "yes"),
    ("scale-corrected ratio primary", str(ref_pct) in MSF, f"{ref_pct}%"),
    ("prespecification described honestly",
     "not** a prespecified confirmatory analysis" in MSF
     and "not a prespecified confirmatory analysis" in re.sub(r"\s+", " ", cov), "yes"),
    ("prior rejection disclosed neutrally",
     "declined by another journal without external review" in re.sub(r"\s+", " ", cov)
     and "Frontiers" not in cov, "neutral wording"),
    ("no outcome model claimed", "estimates no patient outcome" in MSF, "yes"),
    ("cover letter dated", cov.strip().startswith("5 August 2026"), "5 August 2026"),
    ("3 figures present (tif+pdf)",
     all((PKG / f"Figure{i}.{ext}").exists() for i in (1, 2, 3) for ext in ("tif", "pdf")),
     "yes"),
    ("3 documents present",
     all((PKG / n).exists() for n in ("Manuscript.docx", "Cover_Letter.docx",
                                      "Additional_file_1.docx")), "yes"),
]
qa = pd.DataFrame(checks, columns=["check", "pass", "detail"])
qa.to_csv(PKG / "submission_qa.csv", index=False)

(PKG / "SUBMISSION_CHECKLIST.md").write_text("\n".join([
    "# BMJ Quality & Safety submission package", "",
    f"Assembled {datetime.now(timezone.utc).isoformat(timespec='seconds')}", "",
    qa.assign(**{"pass": qa["pass"].map({True: "PASS", False: "FAIL"})}).to_markdown(index=False),
    "", "## Before upload", "",
    "- [x] Code deposited: https://github.com/jiajunluo430-creator/icu-nutrition-interruption-attribution",
    "- [ ] Mint a Zenodo DOI and add it to the data availability statement",
    "- [ ] ORCIDs for all five authors",
    "- [ ] Graphical abstract (BMJ QS recommends one; not generated here)",
    "- [ ] Confirm abstract word count in Word against the 250-word limit",
    "- [ ] Confirm double line spacing if required at submission",
]), encoding="utf-8")

stamp = datetime.now().strftime("%Y-%m-%d")
zp = SUB / f"N2_BMJQS_submission_{stamp}.zip"
with zipfile.ZipFile(zp, "w", zipfile.ZIP_DEFLATED) as z:
    for f in sorted(PKG.rglob("*")):
        if f.is_file():
            z.write(f, f.relative_to(PKG))

print()
print(qa.to_string(index=False))
print(f"\nPASS {int(qa['pass'].sum())}/{len(qa)}")
print(f"package: {zp}  ({zp.stat().st_size/1e6:.2f} MB)")
print(f"sha256 : {hashlib.sha256(zp.read_bytes()).hexdigest()[:32]}")
