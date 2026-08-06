"""N2 step 57 - validate every number in the JCE manuscript, render it, and package it.

Expected values are read from deposited outputs, never typed here. The retraction checks
from review round 6 are carried over so the withdrawn claims cannot reappear.
"""
import hashlib
import json
import re
import shutil
import zipfile
from datetime import datetime, timezone
from pathlib import Path

import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(r"D:\N2_icu_nutrition_delivery_gap")
MAN, SUB, FIG, OUT = (ROOT / "07_manuscript", ROOT / "08_submission",
                      ROOT / "04_figures", ROOT / "03_outputs")
CAN, REV = OUT / "canonical", OUT / "review6"
MS_PATH = MAN / "JCE_main_manuscript.md"

RAW = MS_PATH.read_text(encoding="utf-8")
MS = re.sub(r"\s+", " ", RAW)

C = json.load(open(CAN / "canonical_primary.json"))
CLS = pd.read_csv(CAN / "canonical_class_results.csv").set_index("class")
ERA = pd.read_csv(CAN / "canonical_temporal_transport.csv")
TC = json.load(open(CAN / "canonical_transport_clinical.json"))
R6 = json.load(open(REV / "review6_recompute.json"))
S6 = json.load(open(REV / "review6_supplementary.json"))
SH = json.load(open(REV / "hospital_shrinkage.json"))
LF = json.load(open(REV / "like_for_like_airway.json"))
EI = json.load(open(OUT / "eicu_interface_audit.json"))
EB = json.load(open(OUT / "eicu_background_rate.json"))
EA = json.load(open(OUT / "eicu_ascertainment_diagnostic.json"))
WIN = pd.read_csv(REV / "window_definition_sensitivity.csv")
STR = pd.read_csv(REV / "gap_duration_strata.csv")
SCL = pd.read_csv(REV / "scale_and_priority_free.csv")
SEV = pd.read_csv(REV / "severity_strata.csv")
BUR = TC["burden_distribution"]
AF = R6["attributable_fraction"]["target_only"]
WT = lambda w: WIN[(WIN.window.str.startswith(w))
                   & (WIN.scope.str.startswith("target"))].iloc[0]
ref = SCL[SCL.numerator_scale.str.startswith("reference")].iloc[0]
act = SCL[SCL.numerator_scale.str.startswith("actual")].iloc[0]

checks = []


def chk(name, expected, present=None):
    s = str(expected)
    checks.append({"check": name, "expected": s,
                   "pass": (s in MS) if present is None else bool(present)})


tgt = WT("Span")
for nm, v in [("observed rate", tgt.observed_pct), ("background rate", tgt.background_pct),
              ("excess pp", tgt.excess_pp), ("excess CI", f"{tgt.ci_lo} to {tgt.ci_hi}"),
              ("analysis set", f"{C['n_analysis_set']:,}"),
              ("any-class observed", C["rate"]["obs_pct"]),
              ("any-class background", C["rate"]["null_pct"]),
              ("any-class excess", C["rate"]["excess_pp"]),
              ("null maximum", R6["p_value"]["null_max_pct"]),
              ("causal fraction", AF["causal_fraction_pct"]),
              ("causal fraction CI", f"{AF['ci'][0]} to {AF['ci'][1]}"),
              ("in-gap excess", WT("Procedure inside").excess_pp),
              ("onset+/-1h excess", WT("Onset-centred").excess_pp),
              ("onset-to-2h excess", WT("Gap onset to").excess_pp),
              ("window median", R6["window_length_h"]["median"]),
              ("window max", int(R6["window_length_h"]["max"])),
              ("reference-scale kcal", f"{ref.excess_kcal:,}"),
              ("reference-scale pct", ref.pct_of_shortfall),
              ("actual-scale kcal", f"{act.excess_kcal:,}"),
              ("actual-scale pct", act.pct_of_shortfall),
              ("reference kcal/h", R6["scale_diagnostic"]["reference_kcal_per_h"]),
              ("actual kcal/h", R6["scale_diagnostic"]["actual_pregap_kcal_per_h"]),
              ("scale factor", R6["scale_diagnostic"]["ratio"]),
              ("sensitivity min", f"{C['sensitivity_pct_min']:.2f}"),
              ("sensitivity max", f"{C['sensitivity_pct_max']:.2f}"),
              ("era excess min", ERA.rate_excess_pp.min()),
              ("era excess max", ERA.rate_excess_pp.max()),
              ("era pct min", f"{ERA.pct_of_shortfall.min():.3f}"),
              ("era pct max", f"{ERA.pct_of_shortfall.max():.3f}"),
              ("burden zero", BUR["pct_stays_exactly_zero"]),
              ("burden negative", BUR["pct_stays_negative"]),
              ("top decile share", BUR["top10pct_share_of_gross_positive"]),
              ("tau (1 dp, as rendered in text)", f'{SH["tau_pp"]:.1f}'), ("ICC", SH["icc"]),
              ("pooled rate", SH["pooled_rate_pct"]),
              ("hospitals", SH["hospitals"]),
              ("windows per hospital", f"{SH['median_windows_per_hospital']:,}"),
              ("shrunk p10-p90", f"{SH['shrunk_p10_p90'][0]}% to {SH['shrunk_p10_p90'][1]}%"),
              ("noise p10-p90",
               f"{SH['sampling_noise_only_p10_p90'][0]}% to {SH['sampling_noise_only_p10_p90'][1]}%"),
              ("eICU stays", f"{EB['eicu_stays']:,}"),
              ("ascertainment r", EA["pearson_r"]),
              ("ascertainment variance", f"{100*EA['r_squared']:.0f}%"),
              ("airway eICU rate", LF["eicu_airway_only_pct"]),
              ("airway CI", f"{LF['eicu_airway_only_ci'][0]} to {LF['eicu_airway_only_ci'][1]}"),
              ("airway MIMIC rate", LF["mimic_p1_background_pct"]),
              ("density ratio", LF["density_ratio_mimic_over_eicu"]),
              ("pct stays with airway", LF["eicu_pct_stays_with_airway_event"]),
              ("doc lag median", f"{EI['doc_lag_median_min']:.0f} min"),
              ("doc lag p95", f"{EI['doc_lag_p95_min']:.0f} min"),
              ("doc lag pct", EI["doc_lag_pct_over_60min"]),
              ("doc lag n", f"{EI['doc_lag_n']:,}"),
              ("never fed", f"{S6['m7_cohort']['never_fed']:,}"),
              ("fed one day", f"{S6['m7_cohort']['fed_one_day']:,}"),
              ("pre-init primary", S6["m7_cohort"]["pre_initiation_share_primary_pct"]),
              ("pre-init relaxed", S6["m7_cohort"]["pre_initiation_share_relaxed_pct"]),
              ("same formula", S6["m5_validation"]["same_formula_resumed_pct"]),
              ("same formula+order", S6["m5_validation"]["same_formula_and_order_pct"]),
              ("formula switch", S6["m5_validation"]["formula_switch_pct"]),
              ("midnight artifact", S6["m5_validation"]["midnight_artifact_pct"]),
              ("E-value", S6["m11"]["e_value"]), ("risk ratio", S6["m11"]["risk_ratio"])]:
    chk(nm, v)
chk("pre-onset negative excess", f"\u2212{abs(WT('1 h before').excess_pp)}")
for r in STR.itertuples():
    chk(f"stratum {r.gap_stratum} background", r.background_pct)
chk("severity range", f"{SEV.excess_pp.min()} pp to {SEV.excess_pp.max()} pp",
    f"{SEV.excess_pp.min()}" in MS and f"{SEV.excess_pp.max()}" in MS)
for c in ("P1", "P3", "P2"):
    chk(f"{c} excess", f"+{CLS.loc[c, 'rate_excess_pp']} pp")

BANNED = [("this replicated", "replication claim"),
          ("replicating across", "replication claim"),
          ("two independent databases", "replication claim"),
          ("external validation", "validation claim"),
          ("stand to gain nothing", "patient-targeting claim"),
          ("29-fold", "unstable fold range"),
          ("before any estimate was computed", "overstated prespecification")]
for tok, why in BANNED:
    chk(f"retracted: {why}", f"no '{tok}'", tok not in MS)
chk("no replication claimed", "-", "we make no claim of replication" in MS)
chk("scope limited to timestamp attribution", "-",
    "does not evaluate reasons recorded prospectively at the bedside" in MS)
chk("prespecification honest", "-", "not** a prespecified confirmatory analysis" in MS)
chk("spread not care quality", "-", "must not be read as variation in quality of care" in MS)
chk("nutrition-support cohort", "-", "not an ICU cohort" in MS)
chk("immortal time", "-", "immortal-time selection" in MS)
chk("What is new box", "-", "## What is new?" in RAW)
chk("JCE abstract headings", "-",
    all(h in RAW for h in ("**Objective.**", "**Study Design and Setting.**",
                           "**Results.**", "**Conclusion.**")))
chk("CRediT statement", "-", "CRediT authorship contribution statement" in RAW)
nref = len(re.findall(r"^\[\d+\]\s", RAW.split("## References", 1)[1], re.M))
chk("references", nref, nref >= 20)
chk("Elsevier numbered style", "-", "[1] " in RAW and "doi:" not in
    RAW.split("## References", 1)[1])


def wc(s):
    s = "\n".join(l for l in s.split("\n")
                  if not l.lstrip().startswith("|") and not l.startswith("**Table "))
    return sum(1 for w in re.sub(r"[*`]", "", s).split() if re.search(r"[A-Za-z0-9]", w))


ab = wc(RAW.split("## Abstract", 1)[1].split("**Keywords:**", 1)[0])
bd = wc(RAW.split("## 1. Introduction", 1)[1].split("## Figure legends", 1)[0])
chk("abstract <= 300 words", ab, ab <= 300)
chk("main text <= 4000 words", bd, bd <= 4000)

df = pd.DataFrame(checks)
df.to_csv(OUT / "jce_validation.csv", index=False)
npass = int(df["pass"].sum())
print(f"PASS {npass}/{len(df)}  |  abstract {ab} words, main text {bd:,} words, refs {nref}")
fails = df[~df["pass"]]
if len(fails):
    print("\n=== FAILURES ===")
    print(fails.to_string(index=False))
    raise SystemExit(1)

# ------------------------------------------------------------------ render
_src = (ROOT / "01_scripts" / "12_build_docx.py").read_text(encoding="utf-8")
_ns = {}
exec(compile(_src.split("# ========================================"
                        "======================= manuscript")[0], "h", "exec"), _ns)
base_doc, convert = _ns["base_doc"], _ns["convert"]

AUTHORS = [("Jiajun Luo", "a,b,1"), ("Qinglong Chen", "c,1"), ("Jing Liu", "c,1"),
           ("Fanghui Lu", "c,*"), ("Xiaolong Liang", "a,*")]
AFFIL = ["a  Department of Gastrointestinal Surgery, The First Affiliated Hospital of "
         "Chongqing Medical University, Chongqing, China",
         "b  Molecular Oncology Laboratory, Department of Orthopedic Surgery and "
         "Rehabilitation Medicine, The University of Chicago Medical Center, Chicago, IL, USA",
         "c  Department of Cancer Center, The Second Affiliated Hospital of Chongqing "
         "Medical University, Chongqing, China"]

body = RAW.split("\n", 1)[1].split("<!-- generated back matter -->")[0]
d = base_doc(line_numbers=True)
p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(14)
r = p.add_run(RAW.split("\n", 1)[0].lstrip("# ").strip())
r.bold = True
r.font.size = Pt(15)
p = d.add_paragraph()
p.paragraph_format.space_after = Pt(10)
for i, (nm, sup) in enumerate(AUTHORS):
    if i:
        p.add_run(", ")
    p.add_run(nm)
    p.add_run(sup).font.superscript = True
for a in AFFIL + ["1  These authors contributed equally.",
                  "*  Corresponding authors: 204951@hospital.cqmu.edu.cn (X. Liang), "
                  "lufh@cqmu.edu.cn (F. Lu)"]:
    q = d.add_paragraph()
    q.paragraph_format.space_after = Pt(2)
    q.paragraph_format.line_spacing = 1.15
    q.add_run(a).font.size = Pt(10)
for lbl, val in (("Article type: ", "Original Article"),
                 ("Abstract word count: ", f"{ab:,}"),
                 ("Main text word count: ", f"{bd:,}"),
                 ("Figures: ", "3"), ("Tables: ", "2"), ("References: ", str(nref))):
    q = d.add_paragraph()
    q.paragraph_format.space_after = Pt(2)
    rr = q.add_run(lbl)
    rr.bold = True
    rr.font.size = Pt(10)
    q.add_run(val).font.size = Pt(10)
d.add_page_break()
convert(body, d)
d.save(SUB / "JCE_manuscript.docx")

d = base_doc()
convert(re.sub(r"^# .*$", "", (MAN / "BMJQS_additional_file.md").read_text(encoding="utf-8"),
               flags=re.M), d)
d.save(SUB / "JCE_Supplementary_Material.docx")

d = base_doc()
for s in d.sections:
    s.top_margin = s.bottom_margin = Inches(0.7)
    s.left_margin = s.right_margin = Inches(0.9)
d.styles["Normal"].font.size = Pt(11)
convert((MAN / "cover_letter_jce.md").read_text(encoding="utf-8"), d)
d.save(SUB / "JCE_cover_letter.docx")

PKG = SUB / "jce_package"
if PKG.exists():
    shutil.rmtree(PKG)
PKG.mkdir()
for i, f in enumerate(["figure1_cohort_flow", "figure2_attribution_bmjqs",
                       "figure3_comparability"], 1):
    for ext in ("tif", "pdf"):
        shutil.copy(FIG / f"{f}.{ext}", PKG / f"Fig{i}.{ext}")
for a, b in [("JCE_manuscript.docx", "Manuscript.docx"),
             ("JCE_cover_letter.docx", "Cover_Letter.docx"),
             ("JCE_Supplementary_Material.docx", "Supplementary_Material.docx")]:
    shutil.copy(SUB / a, PKG / b)
for n in ("references_verified.csv", "jce_validation.csv"):
    shutil.copy(OUT / n, PKG / n)
df.to_csv(PKG / "jce_validation.csv", index=False)

stamp = datetime.now().strftime("%Y-%m-%d")
zp = SUB / f"N2_JClinEpi_submission_{stamp}.zip"
with zipfile.ZipFile(zp, "w", zipfile.ZIP_DEFLATED) as z:
    for f in sorted(PKG.rglob("*")):
        if f.is_file():
            z.write(f, f.relative_to(PKG))
print(f"\npackage: {zp}  ({zp.stat().st_size/1e6:.2f} MB)")
print(f"sha256 : {hashlib.sha256(zp.read_bytes()).hexdigest()[:32]}")
for f in sorted(PKG.iterdir()):
    print(f"  {f.name:<32} {f.stat().st_size/1024:>8.1f} KB")
