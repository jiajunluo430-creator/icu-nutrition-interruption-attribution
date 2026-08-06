"""N2 step 39 - render the Critical Care submission set as .docx.

Reuses the converter from script 12 (same markdown subset, same table handling) with a
BMC title page: no line numbers required by Critical Care, but they are kept because they
help reviewers and BMC does not forbid them.
"""
import importlib.util
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(r"D:\N2_icu_nutrition_delivery_gap")
MAN, SUB = ROOT / "07_manuscript", ROOT / "08_submission"
SUB.mkdir(exist_ok=True)

# reuse base_doc/convert/add_runs from script 12 rather than re-implementing them
spec = importlib.util.spec_from_file_location("_d12", ROOT / "01_scripts" / "12_build_docx.py")
_d12 = importlib.util.module_from_spec(spec)
_src = (ROOT / "01_scripts" / "12_build_docx.py").read_text(encoding="utf-8")
_helpers = _src.split("# =============================================================== manuscript")[0]
exec(compile(_helpers, "12_build_docx_helpers", "exec"), _d12.__dict__)
base_doc, convert = _d12.base_doc, _d12.convert

AUTHORS = [("Jiajun Luo", "1,2,†"), ("Qinglong Chen", "3,†"), ("Jing Liu", "3,†"),
           ("Fanghui Lu", "3,*"), ("Xiaolong Liang", "1,*")]
AFFIL = [
    "1  Department of Gastrointestinal Surgery, The First Affiliated Hospital of "
    "Chongqing Medical University, Chongqing, China",
    "2  Molecular Oncology Laboratory, Department of Orthopedic Surgery and "
    "Rehabilitation Medicine, The University of Chicago Medical Center, Chicago, IL, USA",
    "3  Department of Cancer Center, The Second Affiliated Hospital of Chongqing "
    "Medical University, Chongqing, China",
]
CORR = [
    "Xiaolong Liang, MD, PhD \u2014 Department of Gastrointestinal Surgery, The First "
    "Affiliated Hospital of Chongqing Medical University, Chongqing 400016, China. "
    "E-mail: 204951@hospital.cqmu.edu.cn",
    "Fanghui Lu, PhD \u2014 Department of Cancer Center, The Second Affiliated Hospital "
    "of Chongqing Medical University, Chongqing, China. E-mail: lufh@cqmu.edu.cn",
]


def wordcount(s):
    s = re.sub(r"[*`]", "", s)
    return sum(1 for w in s.split() if re.search(r"[A-Za-z0-9]", w))


# =============================================================== manuscript
src = (MAN / "CritCare_main_manuscript.md").read_text(encoding="utf-8")
title = src.split("\n", 1)[0].lstrip("# ").strip()
body = src.split("\n", 1)[1]

d = base_doc(line_numbers=True)

p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_after = Pt(14)
r = p.add_run(title)
r.bold = True
r.font.size = Pt(16)

p = d.add_paragraph()
p.paragraph_format.space_after = Pt(10)
for i, (nm, sup) in enumerate(AUTHORS):
    if i:
        p.add_run(", ")
    p.add_run(nm)
    rs = p.add_run(sup)
    rs.font.superscript = True

for a in AFFIL:
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(a).font.size = Pt(10)

for note in ("\u2020 These authors contributed equally and share first authorship.",
             "* Correspondence:"):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    p.add_run(note).font.size = Pt(10)

for c in CORR:
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Inches(0.25)
    p.add_run(c).font.size = Pt(10)

_body_txt = body.split("## Background", 1)[1].split("## References", 1)[0]
_ab = body.split("## Abstract", 1)[1].split("**Keywords:**", 1)[0]
meta = [("Article type: ", "Research \u2014 Critical Care"),
        ("Abstract word count: ", f"{wordcount(_ab):,}"),
        ("Main text word count (Background to Conclusions): ", f"{wordcount(_body_txt):,}"),
        ("Figures: ", str(len(re.findall(r"^\*\*Fig\. \d", body, re.M)))),
        ("Tables: ", str(len(re.findall(r"^\*\*Table \d", body, re.M)))),
        ("Additional files: ", "1")]
for lbl, val in meta:
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    rr = p.add_run(lbl)
    rr.bold = True
    rr.font.size = Pt(10)
    p.add_run(val).font.size = Pt(10)

d.add_page_break()
convert(body, d)
out = SUB / "CritCare_manuscript.docx"
d.save(out)
print(f"  {out.name}  ({len(d.tables)} tables)")

# =============================================================== cover letter
d = base_doc()
for s in d.sections:
    s.top_margin = s.bottom_margin = Inches(0.7)
    s.left_margin = s.right_margin = Inches(0.9)
d.styles["Normal"].font.size = Pt(11)
convert((MAN / "cover_letter_critcare.md").read_text(encoding="utf-8"), d)
d.save(SUB / "CritCare_cover_letter.docx")
print("  CritCare_cover_letter.docx")

# =============================================================== additional file
d = base_doc()
sup = (MAN / "CritCare_additional_file.md").read_text(encoding="utf-8")
sup = re.sub(r"^# .*$", "", sup, flags=re.M)
sup = re.sub(r"^\*\*(Chance|Background) co-occurrence.*$", "", sup, flags=re.M)
convert(sup, d)
d.save(SUB / "CritCare_Additional_file_1.docx")
print(f"  CritCare_Additional_file_1.docx  ({len(d.tables)} tables)")
print("\nDONE")
